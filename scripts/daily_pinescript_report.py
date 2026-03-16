#!/usr/bin/env python3
"""
Daily PineScript Report — Sends Jean the latest best strategy as a .docx
Runs via cron at 06:00 UTC daily.

Rules:
- Only sends if the best config changed since yesterday
- Includes Python return, est. TV return, MDD, PF
- Word file with paste-ready PineScript
"""
import json
import os
import subprocess
import sys
from datetime import datetime, timezone
from pathlib import Path

LAB_STATE = "/home/yaraclawd/.openclaw/workspace-luna/files/strategy_lab/lab_state.json"
PINE_FILE = "/home/yaraclawd/.openclaw/workspace-luna/files/strategy_lab/gaussian_ls_v7_s_trail.pine"
LAST_SENT = "/home/yaraclawd/.openclaw/workspace-luna/files/strategy_lab/last_sent_pct.txt"
OUTBOUND = "/home/yaraclawd/.openclaw/media/outbound"
CHAT_ID = "7422563444"

def main():
    ts = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")
    print(f"[{ts}] Daily PineScript report starting")

    # Load lab state
    if not os.path.exists(LAB_STATE):
        print("No lab state file — skipping")
        return

    with open(LAB_STATE) as f:
        state = json.load(f)

    best = state.get("best_so_far", {})
    python_pct = best.get("python_pct", 0)
    est_tv = best.get("estimated_tv_pct", 0)
    mdd = best.get("mdd", 0)
    pf = best.get("pf", 0)
    trades = best.get("trades", 0)
    config = best.get("config", "unknown")

    # Check if we already sent this
    last_sent = 0
    if os.path.exists(LAST_SENT):
        with open(LAST_SENT) as f:
            try:
                last_sent = float(f.read().strip())
            except:
                last_sent = 0

    if python_pct <= last_sent:
        print(f"No improvement: current {python_pct}% <= last sent {last_sent}%. Skipping.")
        return

    # Check PineScript exists
    if not os.path.exists(PINE_FILE):
        print(f"PineScript not found at {PINE_FILE} — skipping")
        return

    # Build Word doc
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError:
        print("python-docx not installed — skipping")
        return

    doc = Document()
    today = datetime.now(timezone.utc).strftime("%Y-%m-%d")
    doc.add_heading(f'Gaussian L+S v7 — Daily Update ({today})', level=1)

    doc.add_heading('Performance', level=2)
    table = doc.add_table(rows=6, cols=2, style='Table Grid')
    data = [
        ('Config', config),
        ('Python Return (net)', f'+{python_pct:,.0f}%'),
        ('Est. TV Return', f'~+{est_tv:,.0f}%'),
        ('Max Drawdown', f'{mdd:.1f}%'),
        ('Profit Factor', f'{pf:.2f}'),
        ('Trades', str(trades)),
    ]
    for i, (k, v) in enumerate(data):
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = v

    doc.add_paragraph('')
    doc.add_heading('Instructions', level=2)
    doc.add_paragraph('1. TradingView → Pine Editor → paste code below')
    doc.add_paragraph('2. Chart: BTCUSDT.P Binance, 1D, from 2018-01-01')
    doc.add_paragraph('3. Properties: Commission 0.055%, Initial Capital 1000')

    doc.add_paragraph('')
    doc.add_heading('PineScript Code', level=2)

    with open(PINE_FILE) as f:
        code = f.read()

    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = 'Courier New'
    run.font.size = Pt(8)

    os.makedirs(OUTBOUND, exist_ok=True)
    docx_path = os.path.join(OUTBOUND, f"Gaussian_v7_{today}.docx")
    doc.save(docx_path)
    print(f"Word doc saved: {docx_path}")

    # Send via openclaw
    msg = (
        f"📊 Daily Strategy Update — {today}\n\n"
        f"Config: {config}\n"
        f"Python: +{python_pct:,.0f}% | MDD {mdd:.1f}% | PF {pf:.2f}\n"
        f"Est TV: ~+{est_tv:,.0f}% (benchmark: v6 +1,781%)\n\n"
        f"PineScript attached — paste into TV Pine Editor."
    )

    result = subprocess.run(
        ["openclaw", "message", "send",
         "--channel", "telegram", "--account", "luna",
         "--target", CHAT_ID,
         "--media", docx_path,
         "-m", msg],
        capture_output=True, text=True, timeout=30
    )

    if result.returncode == 0:
        print(f"Sent to Jean: +{python_pct}% (~{est_tv}% TV)")
        # Record what we sent
        with open(LAST_SENT, "w") as f:
            f.write(str(python_pct))
    else:
        print(f"Send failed: {result.stderr}")

if __name__ == "__main__":
    main()
